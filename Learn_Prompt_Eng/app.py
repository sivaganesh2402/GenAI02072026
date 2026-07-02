import os
import re
import uuid
from pathlib import Path
from typing import Dict, List, Optional

from flask import Flask, flash, redirect, render_template, request, send_file, url_for
from docx import Document
from PyPDF2 import PdfReader

from generate_test_plan import write_test_plan_docx

app = Flask(__name__)
app.secret_key = os.environ.get('FLASK_SECRET', 'change-me-please')

ALLOWED_EXTENSIONS = {'.docx', '.pdf', '.md'}
OUTPUT_FORMATS = {'docx'}
UPLOAD_FOLDER = Path(__file__).parent / 'uploads'
OUTPUT_FOLDER = Path(__file__).parent / 'outputs'
UPLOAD_FOLDER.mkdir(parents=True, exist_ok=True)
OUTPUT_FOLDER.mkdir(parents=True, exist_ok=True)

SECTION_KEYWORDS = {
    'Executive Summary': ['executive summary', 'summary'],
    'Product Vision': ['product vision'],
    'Target Users': ['target users', 'primary users', 'secondary users', 'user base'],
    'Business Objectives': ['business objectives', 'objectives'],
    'Current State Analysis': ['current state analysis', 'existing features'],
    'Functional Requirements': ['functional requirements', 'authentication system', 'password management', 'user input validation', 'user experience features', 'interface design', 'accessibility features', 'branding and visual design'],
    'Security Requirements': ['security specifications', 'data protection', 'compliance standards', 'rate limiting', 'secure storage', 'https enforcement'],
    'Performance Requirements': ['performance requirements', 'load time optimization', 'scalability', 'high availability', 'concurrent users', 'geographic distribution'],
    'Integration Requirements': ['integration requirements', 'platform integrations', 'third-party services', 'analytics integration', 'customer support', 'enterprise sso', 'social login', 'marketing tools'],
    'User Journey and Flow': ['user journey and flow', 'new user experience', 'returning user experience', 'error recovery flow'],
    'Success Metrics and KPIs': ['success metrics and kpis', 'performance metrics', 'security metrics', 'business metrics'],
    'Implementation Considerations': ['implementation considerations', 'development phases'],
    'Risk Mitigation': ['risk mitigation', 'security risks', 'performance risks'],
    'Future Enhancements': ['future enhancements', 'advanced features'],
}


def allowed_file(filename: str) -> bool:
    return Path(filename).suffix.lower() in ALLOWED_EXTENSIONS


def clean_line(text: str) -> str:
    text = text.strip()
    text = re.sub(r'\s*\d+$', '', text)
    return text


def normalize_heading(text: str) -> str:
    return text.lower().strip(':-. ')


def guess_section_title(line: str) -> Optional[str]:
    text = normalize_heading(line)
    for section, keywords in SECTION_KEYWORDS.items():
        for keyword in keywords:
            if keyword in text:
                return section
    return None


def extract_lines_from_docx(path: Path) -> List[str]:
    doc = Document(path)
    lines: List[str] = []
    for paragraph in doc.paragraphs:
        text = clean_line(paragraph.text)
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(clean_line(cell.text) for cell in row.cells if clean_line(cell.text))
            if row_text:
                lines.append(row_text)
    return lines


def extract_lines_from_md(path: Path) -> List[str]:
    with open(path, 'r', encoding='utf-8') as f:
        return [clean_line(line) for line in f if clean_line(line)]


def extract_lines_from_pdf(path: Path) -> List[str]:
    reader = PdfReader(str(path))
    text = []
    for page in reader.pages:
        page_text = page.extract_text() or ''
        page_lines = [clean_line(line) for line in page_text.splitlines() if clean_line(line)]
        text.extend(page_lines)
    return text


def build_section_map(lines: List[str]) -> Dict[str, List[str]]:
    sections: Dict[str, List[str]] = {}
    current_section: Optional[str] = None
    document_title: Optional[str] = None

    for index, line in enumerate(lines):
        if index == 0 and 'product requirements document' in line.lower():
            document_title = line
            continue

        section = guess_section_title(line)
        if section:
            current_section = section
            sections.setdefault(current_section, [])
            continue

        if current_section is None:
            document_title = document_title or line
            continue

        sections.setdefault(current_section, []).append(line)

    if document_title and 'Document Overview' not in sections:
        sections.setdefault('Document Overview', [])
        if document_title not in sections['Document Overview']:
            sections['Document Overview'].append(document_title)

    return sections


def extract_bullets(lines: List[str]) -> List[str]:
    bullets: List[str] = []
    for line in lines:
        if line.startswith(('-', '*')):
            bullets.append(line.lstrip('-* ').strip())
        elif re.match(r'^\d+[\.)]\s+', line):
            bullets.append(re.sub(r'^\d+[\.)]\s+', '', line))
        elif ':' in line and len(line.split(':', 1)[0]) < 30 and line.split(':', 1)[1].strip():
            bullets.append(line)
        else:
            if bullets and not bullets[-1].endswith(('.', '?', '!')):
                bullets[-1] = bullets[-1] + ' ' + line
            else:
                bullets.append(line)
    return [clean_line(b) for b in bullets if b]


def title_from_sections(sections: Dict[str, List[str]]) -> str:
    return sections.get('Document Overview', ['Generated Test Plan'])[0]


def build_intro(sections: Dict[str, List[str]]) -> List[str]:
    intro: List[str] = []
    title = title_from_sections(sections)
    intro.append(f'Document: {title}')

    if summary := sections.get('Executive Summary'):
        intro.append('Executive Summary:')
        intro.extend(extract_bullets(summary))

    if vision := sections.get('Product Vision'):
        intro.append('Product Vision:')
        intro.extend(extract_bullets(vision))

    if objectives := sections.get('Business Objectives'):
        intro.append('Business Objectives:')
        intro.extend(extract_bullets(objectives))

    if not intro:
        intro.append('Insufficient information to determine — requires clarification.')
    return intro


def gather_feature_statements(sections: Dict[str, List[str]]) -> List[str]:
    feature_sections = [
        'Functional Requirements',
        'Current State Analysis',
        'User Journey and Flow',
        'Integration Requirements',
        'Security Requirements',
        'Performance Requirements',
    ]
    features: List[str] = []
    for key in feature_sections:
        if key in sections:
            features.extend(extract_bullets(sections[key]))
    return features


def build_scope(sections: Dict[str, List[str]]) -> List[str]:
    in_scope = gather_feature_statements(sections)
    if not in_scope:
        in_scope = ['Insufficient information to determine — requires clarification.']
    else:
        in_scope = ['The following capabilities are within scope:'] + in_scope

    out_of_scope = [
        'Non-login modules of the VWO platform that are outside the authentication and entry-point workflow.',
        'Any dashboard or experimentation features beyond successful authentication and transition to the main platform.',
    ]
    return in_scope + [''] + ['The following capabilities are considered out of scope:'] + out_of_scope


def build_strategy(sections: Dict[str, List[str]]) -> List[str]:
    return [
        'Functional testing of authentication and login workflow, including password reset, optional 2FA, and SSO flows.',
        'Usability and accessibility testing for responsive design, keyboard navigation, ARIA labels, and high contrast modes.',
        'Security testing for session management, HTTPS enforcement, encrypted credentials, rate limiting, and enterprise access control.',
        'Performance testing for page load, concurrent login attempts, and global availability targets.',
        'Regression testing for all core login and recovery workflows after each build update.',
    ]


def build_features(sections: Dict[str, List[str]]) -> List[str]:
    features = gather_feature_statements(sections)
    if not features:
        return ['Insufficient information to determine — requires clarification.']
    return features


def build_environment(sections: Dict[str, List[str]]) -> List[str]:
    return [
        'Target environments are inferred from the PRD focus on responsive login experience and global enterprise users.',
        'Desktop browsers: Chrome, Firefox, Edge, Safari.',
        'Mobile form factors: Android Chrome and iOS Safari for touch-friendly login handling.',
        'Network conditions: standard broadband and mobile network emulation for performance verification.',
        'Authentication test data should include valid/invalid credentials, account recovery flows, and enterprise SSO scenarios.',
    ]


def build_entry_exit(_: Dict[str, List[str]]) -> Dict[str, List[str]]:
    return {
        'entry': [
            'PRD is attached and verified by QA.',
            'Test environment for the login dashboard is provisioned and accessible.',
            'Required user credentials and enterprise login test accounts are available.',
            'A stable build containing the login dashboard functionality is deployed for testing.',
        ],
        'exit': [
            'All critical login, password recovery, and SSO scenarios execute successfully or have documented defects.',
            'Security and performance checks defined in the PRD are completed.',
            'Test cases are executed and associated defects are logged with reproducible steps.',
            'Test summary and defect reports are reviewed by the QA lead.',
        ],
    }


def build_roles(_: Dict[str, List[str]]) -> List[str]:
    return [
        'QA Lead: Own the test plan, validate test coverage, and track exit criteria.',
        'Test Engineers: Execute functional, usability, security, and performance tests.',
        'Development Team: Address defects and provide environment support.',
        'Product Owner / Stakeholders: Review test execution results and approve sign-off.',
    ]


def build_deliverables(_: Dict[str, List[str]]) -> List[str]:
    return [
        'Test Plan document',
        'Detailed test cases and execution status',
        'Defect reports with severity and reproduction steps',
        'Test summary and sign-off report',
    ]


def build_risks(sections: Dict[str, List[str]]) -> List[str]:
    if mitigation := sections.get('Risk Mitigation'):
        bullets = extract_bullets(mitigation)
    else:
        bullets = []
    return bullets or ['Insufficient information to determine — requires clarification.']


def build_schedule(_: Dict[str, List[str]]) -> List[str]:
    return ['Insufficient information to determine — requires clarification.']


def build_assumptions(_: Dict[str, List[str]]) -> List[str]:
    return [
        'The PRD scope is limited to the VWO login dashboard and authentication entry point.',
        'Non-login platform functionality is not included unless explicitly stated in the PRD.',
        'Accessibility and responsive requirements apply to the login dashboard interface.',
    ]


def build_approval(_: Dict[str, List[str]]) -> List[str]:
    return ['Insufficient information to determine — requires clarification.']


def add_section(doc: Document, title: str, lines: List[str]) -> None:
    doc.add_heading(title, level=1)
    if not lines:
        doc.add_paragraph('Insufficient information to determine — requires clarification.')
        return
    for line in lines:
        if line == '':
            doc.add_paragraph('')
        elif line.startswith('The following') or line.endswith(':') or line.startswith('Document:'):
            doc.add_paragraph(line)
        else:
            doc.add_paragraph(line, style='List Bullet')


def render_markdown(sections: Dict[str, List[str]]) -> str:
    parts = [
        '# Test Plan',
        '## 1. Introduction & Objectives',
    ]
    for line in build_intro(sections):
        parts.append(f'- {line}' if not line.endswith(':') else line)
    parts.append('## 2. Scope (In-Scope / Out-of-Scope)')
    for line in build_scope(sections):
        parts.append(f'- {line}' if not line.endswith(':') else line)
    parts.append('## 3. Test Strategy')
    for line in build_strategy(sections):
        parts.append(f'- {line}')
    parts.append('## 4. Features to be Tested / Not to be Tested')
    for line in build_features(sections):
        parts.append(f'- {line}')
    parts.append('## 5. Test Environment & Test Data Requirements')
    for line in build_environment(sections):
        parts.append(f'- {line}')
    entry_exit = build_entry_exit(sections)
    parts.append('## 6. Entry Criteria')
    for line in entry_exit['entry']:
        parts.append(f'- {line}')
    parts.append('## 7. Exit Criteria')
    for line in entry_exit['exit']:
        parts.append(f'- {line}')
    parts.append('## 8. Roles & Responsibilities')
    for line in build_roles(sections):
        parts.append(f'- {line}')
    parts.append('## 9. Test Deliverables')
    for line in build_deliverables(sections):
        parts.append(f'- {line}')
    parts.append('## 10. Risks & Mitigations')
    for line in build_risks(sections):
        parts.append(f'- {line}')
    parts.append('## 11. Schedule / Milestones')
    for line in build_schedule(sections):
        parts.append(f'- {line}')
    parts.append('## 12. Assumptions & Dependencies')
    for line in build_assumptions(sections):
        parts.append(f'- {line}')
    parts.append('## 13. Approval / Sign-off')
    for line in build_approval(sections):
        parts.append(f'- {line}')
    return '\n'.join(parts)


def write_docx(text: str, output_path: Path) -> None:
    doc = Document()
    doc.add_heading('Test Plan', level=0)
    for line in text.splitlines():
        if not line.strip():
            doc.add_paragraph('')
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)
    doc.save(output_path)


def process_prd_file(path: Path) -> Dict[str, List[str]]:
    suffix = path.suffix.lower()
    if suffix == '.docx':
        lines = extract_lines_from_docx(path)
    elif suffix == '.md':
        lines = extract_lines_from_md(path)
    elif suffix == '.pdf':
        lines = extract_lines_from_pdf(path)
    else:
        raise ValueError('Unsupported file type')
    return build_section_map(lines)


def create_output_files(sections: Dict[str, List[str]], output_basename: str) -> Dict[str, Path]:
    paths: Dict[str, Path] = {}

    docx_path = OUTPUT_FOLDER / f'{output_basename}.docx'
    write_test_plan_docx(sections, docx_path)
    paths['docx'] = docx_path

    return paths


def cleanup_old_files(folder: Path, keep: Optional[List[str]] = None) -> None:
    keep = keep or []
    for path in folder.iterdir():
        if path.name not in keep and path.is_file():
            try:
                path.unlink()
            except OSError:
                pass


def create_unique_filename(filename: str) -> str:
    base = Path(filename).stem
    unique = f'{base}-{uuid.uuid4().hex[:8]}'
    return unique


@app.route('/', methods=['GET', 'POST'])
def index():
    output_files = None
    upload_status = None
    generation_status = None
    if request.method == 'POST':
        cleanup_old_files(UPLOAD_FOLDER)
        cleanup_old_files(OUTPUT_FOLDER)

        uploaded_file = request.files.get('prd_file')
        if not uploaded_file or uploaded_file.filename == '':
            flash('Please upload a PRD file.', 'danger')
            return redirect(request.url)

        if not allowed_file(uploaded_file.filename):
            flash('Unsupported file type. Use DOCX, PDF, or MD.', 'danger')
            return redirect(request.url)

        saved_name = create_unique_filename(uploaded_file.filename)
        upload_path = UPLOAD_FOLDER / f'{saved_name}{Path(uploaded_file.filename).suffix.lower()}'
        uploaded_file.save(upload_path)
        upload_status = f'File uploaded successfully: {uploaded_file.filename}'

        try:
            sections = process_prd_file(upload_path)
            output_basename = create_unique_filename('test_plan')
            output_paths = create_output_files(sections, output_basename)
            output_files = output_paths
            if output_paths.get('docx') and output_paths['docx'].exists():
                generation_status = 'Test plan generated and ready to download.'
            else:
                generation_status = None
                flash('The DOCX file was not created. Please try again.', 'danger')
        except Exception as exc:
            flash(f'Error processing file: {exc}', 'danger')

    return render_template(
        'index.html',
        output_files=output_files,
        upload_status=upload_status,
        generation_status=generation_status,
    )


@app.route('/download/<format>/<filename>')
def download(format: str, filename: str):
    if format not in OUTPUT_FORMATS:
        flash('Invalid download format.', 'danger')
        return redirect(url_for('index'))

    file_path = OUTPUT_FOLDER / f'{filename}.{format}'
    if not file_path.exists():
        flash('Requested file not found.', 'danger')
        return redirect(url_for('index'))

    return send_file(str(file_path), as_attachment=True)


if __name__ == '__main__':
    app.run(debug=True)
