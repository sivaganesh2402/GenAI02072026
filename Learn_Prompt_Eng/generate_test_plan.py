import argparse
import re
from datetime import date
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SECTION_KEYWORDS = {
    'Executive Summary': ['executive summary', 'summary'],
    'Product Vision': ['product vision'],
    'Target Users': ['target users', 'primary users', 'secondary users', 'user base'],
    'Business Objectives': ['business objectives', 'objectives'],
    'Current State Analysis': ['current state analysis', 'existing features'],
    'Functional Requirements': [
        'functional requirements',
        'authentication system',
        'password management',
        'user input validation',
        'user experience features',
        'interface design',
        'accessibility features',
        'branding and visual design',
    ],
    'Security Requirements': [
        'security specifications',
        'data protection',
        'compliance standards',
        'rate limiting',
        'secure storage',
        'https enforcement',
    ],
    'Performance Requirements': [
        'performance requirements',
        'load time optimization',
        'scalability',
        'high availability',
        'concurrent users',
        'geographic distribution',
    ],
    'Integration Requirements': [
        'integration requirements',
        'platform integrations',
        'third-party services',
        'analytics integration',
        'customer support',
        'enterprise sso',
        'social login',
        'marketing tools',
    ],
    'User Journey and Flow': ['user journey and flow', 'new user experience', 'returning user experience', 'error recovery flow'],
    'Success Metrics and KPIs': ['success metrics and kpis', 'performance metrics', 'security metrics', 'business metrics'],
    'Implementation Considerations': ['implementation considerations', 'development phases'],
    'Risk Mitigation': ['risk mitigation', 'security risks', 'performance risks'],
    'Future Enhancements': ['future enhancements', 'advanced features'],
}

TEST_PLAN_SECTIONS = [
    'Objective',
    'Scope',
    'Inclusions',
    'Exclusions',
    'Test Environments',
    'Defect Reporting Procedure',
    'Test Strategy',
    'Test Schedule',
    'Test Deliverables',
    'Entry and Exit Criteria',
    'Test Execution',
    'Test Closure',
    'Tools',
    'Risks and Mitigations',
    'Approvals',
]

ACCENT = RGBColor(31, 78, 121)
LIGHT_FILL = 'D9EAF7'
BORDER = 'B7C9D6'
BODY = RGBColor(31, 31, 31)
GAP = 'Insufficient information to determine.'


def clean_line(text: str) -> str:
    text = text.strip()
    text = re.sub(r'\s*\d+$', '', text)
    return text


def normalize_heading(text: str) -> str:
    text = re.sub(r'^\d+[\.)]\s*', '', text)
    return text.lower().strip(':-. ')


def guess_section_title(line: str) -> Optional[str]:
    text = normalize_heading(line)
    for section, keywords in SECTION_KEYWORDS.items():
        for keyword in keywords:
            if keyword in text:
                return section
    return None


def extract_lines_from_docx(path: Path) -> List[str]:
    doc = Document(path)
    lines: List[str] = []
    for paragraph in doc.paragraphs:
        text = clean_line(paragraph.text)
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(clean_line(cell.text) for cell in row.cells if clean_line(cell.text))
            if row_text:
                lines.append(row_text)
    return lines


def build_section_map(lines: List[str]) -> Dict[str, List[str]]:
    sections: Dict[str, List[str]] = {}
    current_section: Optional[str] = None
    document_title: Optional[str] = None

    for index, line in enumerate(lines):
        if index == 0 and 'product requirements document' in line.lower():
            document_title = line
            continue

        section = guess_section_title(line)
        if section:
            current_section = section
            sections.setdefault(current_section, [])
            continue

        if current_section is None:
            document_title = document_title or line
            continue

        sections.setdefault(current_section, []).append(line)

    if document_title and 'Document Overview' not in sections:
        sections.setdefault('Document Overview', [])
        if document_title not in sections['Document Overview']:
            sections['Document Overview'].append(document_title)

    return sections


def extract_bullets(lines: Iterable[str]) -> List[str]:
    bullets: List[str] = []
    for raw_line in lines:
        line = clean_line(raw_line)
        if not line:
            continue
        if line.startswith(('-', '*', '•', '▪')):
            bullets.append(line.lstrip('-*•▪ ').strip())
        elif re.match(r'^\d+[\.)]\s+', line):
            bullets.append(re.sub(r'^\d+[\.)]\s+', '', line))
        elif ':' in line and len(line.split(':', 1)[0]) < 35 and line.split(':', 1)[1].strip():
            bullets.append(line)
        elif bullets and not bullets[-1].endswith(('.', '?', '!', ':')):
            bullets[-1] = f'{bullets[-1]} {line}'
        else:
            bullets.append(line)
    return [clean_line(b) for b in bullets if clean_line(b)]


def title_from_sections(sections: Dict[str, List[str]]) -> str:
    title_candidates = [s for s in sections.get('Document Overview', []) if s]
    if title_candidates:
        title = title_candidates[0]
        title = re.sub(r'product requirements document\s*[-:]*\s*', '', title, flags=re.IGNORECASE).strip()
        return title or 'Generated Test Plan'
    return 'Generated Test Plan'


def product_name(sections: Dict[str, List[str]]) -> str:
    title = title_from_sections(sections)
    title = re.sub(r'\b(PRD|Requirements?|Document|Product)\b', '', title, flags=re.IGNORECASE)
    title = re.sub(r'\s+', ' ', title).strip(' -:')
    return title or 'Application Under Test'


def gather_feature_statements(sections: Dict[str, List[str]]) -> List[str]:
    feature_sections = [
        'Functional Requirements',
        'Current State Analysis',
        'User Journey and Flow',
        'Integration Requirements',
        'Security Requirements',
        'Performance Requirements',
        'Future Enhancements',
    ]
    features: List[str] = []
    for key in feature_sections:
        features.extend(extract_bullets(sections.get(key, [])))
    return dedupe(features)


def gather_feature_statements_with_source(sections: Dict[str, List[str]]) -> List[str]:
    feature_sections = [
        'Functional Requirements',
        'Current State Analysis',
        'User Journey and Flow',
        'Integration Requirements',
        'Security Requirements',
        'Performance Requirements',
        'Future Enhancements',
    ]
    features: List[str] = []
    for key in feature_sections:
        for item in extract_bullets(sections.get(key, [])):
            features.append(f'{item} (Source: {key})')
    return dedupe(features)


def source_backed_test_types(sections: Dict[str, List[str]]) -> List[str]:
    mappings = [
        ('Functional Requirements', 'Functional Testing'),
        ('Current State Analysis', 'Regression Testing'),
        ('User Journey and Flow', 'End-to-End Flow Testing'),
        ('Integration Requirements', 'Integration Testing'),
        ('Security Requirements', 'Security Testing'),
        ('Performance Requirements', 'Performance Testing'),
        ('Target Users', 'Usability and Compatibility Review'),
    ]
    scope: List[str] = []
    for section_name, test_type in mappings:
        if sections.get(section_name):
            scope.append(
                f'Inference (low confidence): {test_type} is included because the source contains {section_name}.'
            )
    return scope or [GAP]


def dedupe(items: Iterable[str], limit: Optional[int] = None) -> List[str]:
    seen = set()
    result = []
    for item in items:
        key = re.sub(r'\s+', ' ', item.lower()).strip()
        if not key or key in seen:
            continue
        seen.add(key)
        result.append(item)
        if limit and len(result) >= limit:
            break
    return result


def summarize_from_sections(sections: Dict[str, List[str]], keys: List[str], fallback: str, limit: int = 5) -> List[str]:
    values: List[str] = []
    for key in keys:
        values.extend(f'{item} (Source: {key})' for item in extract_bullets(sections.get(key, [])))
    return dedupe(values, limit=limit) or [fallback]


def build_test_plan_model(sections: Dict[str, List[str]]) -> Dict[str, object]:
    name = product_name(sections)
    feature_statements = gather_feature_statements(sections)
    objective_source = summarize_from_sections(
        sections,
        ['Executive Summary', 'Product Vision', 'Business Objectives'],
        GAP,
        limit=4,
    )
    risk_items = summarize_from_sections(
        sections,
        ['Risk Mitigation', 'Security Requirements', 'Performance Requirements'],
        GAP,
        limit=5,
    )

    inclusions = gather_feature_statements_with_source(sections)[:12] or [GAP]
    exclusions = [GAP]
    scope_categories = source_backed_test_types(sections)

    return {
        'name': name,
        'source_title': title_from_sections(sections),
        'objective': objective_source,
        'scope': scope_categories,
        'inclusions': inclusions,
        'exclusions': exclusions,
        'environment_rows': [
            ('QA', GAP),
            ('Pre Prod', GAP),
            ('Production', GAP),
        ],
        'defect_rows': [
            ('Defect Tool / Process', GAP),
            ('QA Point of Contact', GAP),
            ('Frontend Owner', GAP),
            ('Backend Owner', GAP),
            ('DevOps / Environment Owner', GAP),
        ],
        'strategy': [
            GAP if not feature_statements else f'Create test scenarios from the source-backed requirements listed in Inclusions. (Source: uploaded requirements document)',
        ],
        'schedule_rows': [
            ('Creating Test Plan', 'TBD'),
            ('Test Case Creation', 'TBD'),
            ('Test Case Execution', 'TBD'),
            ('Defect Retest and Regression', 'TBD'),
            ('Summary Reports Submission Date', 'TBD'),
        ],
        'deliverables': [
            'Test Plan document (requested output)',
            GAP,
        ],
        'entry_exit': {
            'Entry Criteria': [
                'Uploaded source document is available to the generator.',
                GAP,
            ],
            'Exit Criteria': [
                GAP,
            ],
        },
        'execution_entry_exit': {
            'Entry Criteria': [
                GAP,
            ],
            'Exit Criteria': [
                GAP,
            ],
        },
        'closure_entry_exit': {
            'Entry Criteria': [
                GAP,
            ],
            'Exit Criteria': [
                GAP,
            ],
        },
        'tools': [
            GAP,
        ],
        'risks': risk_items,
        'approvals': [
            ('QA Lead', GAP),
            ('Product Owner', GAP),
            ('Development Lead', GAP),
            ('Business Stakeholder', GAP),
        ],
        'anti_hallucination': [
            'Grounding rule applied: source-backed items include a source section label where available.',
            f'Missing or unclear details are marked exactly as "{GAP}"',
            'Low-confidence inferred test types are explicitly labeled as inferences.',
        ],
    }


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ''
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.font.color.rgb = BODY


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for margin_name, margin_value in {'top': top, 'start': start, 'bottom': bottom, 'end': end}.items():
        node = tc_mar.find(qn(f'w:{margin_name}'))
        if node is None:
            node = OxmlElement(f'w:{margin_name}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(margin_value))
        node.set(qn('w:type'), 'dxa')


def set_cell_border(cell, color: str = BORDER, size: str = '6') -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in('w:tcBorders')
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = f'w:{edge}'
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), size)
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), color)


def style_table(table, header_fill: str = LIGHT_FILL) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)
            if row_index == 0:
                shade_cell(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = ACCENT


def set_column_widths(table, widths: List[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)


def add_table(doc: Document, headers: Tuple[str, str], rows: List[Tuple[str, str]], widths: Tuple[float, float] = (2.1, 4.2)):
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    set_cell_text(table.rows[0].cells[0], headers[0], bold=True)
    set_cell_text(table.rows[0].cells[1], headers[1], bold=True)
    for left, right in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], left)
        set_cell_text(cells[1], right)
    set_column_widths(table, list(widths))
    style_table(table)
    return table


def add_bullets(doc: Document, items: Iterable[str], limit: Optional[int] = None) -> None:
    for item in dedupe(items, limit=limit):
        paragraph = doc.add_paragraph(style='List Bullet')
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(item)
        run.font.name = 'Arial'
        run.font.size = Pt(10.5)


def add_numbered_items(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style='List Number')
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(item)
        run.font.name = 'Arial'
        run.font.size = Pt(10.5)


def add_heading(doc: Document, text: str, level: int = 2) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
    paragraph.paragraph_format.space_after = Pt(5)
    for run in paragraph.runs:
        run.font.name = 'Arial'
        run.font.color.rgb = ACCENT


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10.5)
    run.font.color.rgb = BODY


def configure_document(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    header = section.header.paragraphs[0]
    header.text = f'Test Plan - {title}'
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('Confidential - QA Test Plan')
    run.font.name = 'Arial'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 100, 100)


def add_cover(doc: Document, model: Dict[str, object]) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run('TEST PLAN')
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(24)
    run.font.color.rgb = ACCENT

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    sub_run = subtitle.add_run(str(model['name']))
    sub_run.bold = True
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = BODY

    metadata_rows = [
        ('Document Type', 'QA Test Plan'),
        ('Project / Application', str(model['name'])),
        ('Source Document', str(model['source_title'])),
        ('Version', '1.0'),
        ('Prepared Date', date.today().strftime('%d %b %Y')),
        ('Status', 'Draft'),
    ]
    add_table(doc, ('Field', 'Details'), metadata_rows, widths=(2.0, 4.5))
    doc.add_paragraph('')


def add_page_index(doc: Document) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run('Page Index')
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(18)
    run.font.color.rgb = ACCENT

    toc_paragraph = doc.add_paragraph()
    fld_simple = OxmlElement('w:fldSimple')
    fld_simple.set(qn('w:instr'), 'TOC \\o "2-4" \\h \\z \\u')
    run_element = OxmlElement('w:r')
    text_element = OxmlElement('w:t')
    text_element.text = 'Right-click and update field to refresh page numbers.'
    run_element.append(text_element)
    fld_simple.append(run_element)
    toc_paragraph._p.append(fld_simple)


def enable_field_update_on_open(doc: Document) -> None:
    settings = doc.settings.element
    update_fields = settings.find(qn('w:updateFields'))
    if update_fields is None:
        update_fields = OxmlElement('w:updateFields')
        settings.append(update_fields)
    update_fields.set(qn('w:val'), 'true')


def add_entry_exit_group(doc: Document, title: str, data: Dict[str, List[str]]) -> None:
    add_heading(doc, title, level=3)
    for subheading, items in data.items():
        add_heading(doc, subheading + ':', level=4)
        add_bullets(doc, items)


def write_test_plan_docx(sections: Dict[str, List[str]], output_path: Path) -> None:
    model = build_test_plan_model(sections)
    doc = Document()
    configure_document(doc, str(model['name']))
    add_cover(doc, model)
    doc.add_page_break()
    add_page_index(doc)
    doc.add_page_break()

    add_heading(doc, 'Objective', level=2)
    add_body_paragraph(
        doc,
        f"This test plan is generated from the uploaded source document for {model['name']}. "
        f'Any detail not present in the source is marked as "{GAP}".',
    )
    add_bullets(doc, model['objective'], limit=4)

    add_heading(doc, 'Scope', level=2)
    add_body_paragraph(doc, f'Scope of Test Plan for {model["name"]}:')
    add_numbered_items(doc, model['scope'])

    add_heading(doc, 'Inclusions', level=3)
    add_bullets(doc, model['inclusions'])

    add_heading(doc, 'Exclusions', level=3)
    add_bullets(doc, model['exclusions'])

    add_heading(doc, 'Test Environments', level=3)
    add_table(doc, ('Name', 'Environment URL / Notes'), model['environment_rows'])

    add_heading(doc, 'Defect Reporting Procedure', level=3)
    add_body_paragraph(
        doc,
        f'Defect workflow details are included only when they are present in the source. Otherwise, they are marked as "{GAP}".',
    )
    add_table(doc, ('Defect Process', 'POC'), model['defect_rows'])

    add_heading(doc, 'Test Strategy', level=3)
    add_bullets(doc, model['strategy'])

    add_heading(doc, 'Test Schedule', level=3)
    add_table(doc, ('Task', 'Dates'), model['schedule_rows'])

    add_heading(doc, 'Test Deliverables.', level=3)
    add_bullets(doc, model['deliverables'])

    add_entry_exit_group(doc, 'Entry and Exit Criteria', model['entry_exit'])
    add_entry_exit_group(doc, 'Test Execution', model['execution_entry_exit'])
    add_entry_exit_group(doc, 'Test Closure', model['closure_entry_exit'])

    add_heading(doc, 'Tools', level=4)
    add_bullets(doc, model['tools'])

    add_heading(doc, 'Risks and Mitigations', level=4)
    add_bullets(doc, model['risks'])

    add_heading(doc, 'Anti-Hallucination Check', level=4)
    add_bullets(doc, model['anti_hallucination'])

    add_heading(doc, 'Approvals', level=4)
    add_table(doc, ('Approver', 'Signature / Date'), model['approvals'])

    enable_field_update_on_open(doc)
    doc.save(output_path)


def render_markdown(sections: Dict[str, List[str]]) -> str:
    model = build_test_plan_model(sections)
    lines = [
        '# TEST PLAN',
        f'## {model["name"]}',
        '',
        '## Objective',
        f'The objective of this test plan is to verify {model["name"]} against the supplied requirements.',
    ]
    for item in model['objective']:
        lines.append(f'- {item}')

    for section_name in TEST_PLAN_SECTIONS[1:]:
        lines.extend(['', f'## {section_name}'])
        if section_name == 'Scope':
            for item in model['scope']:
                lines.append(f'1. {item}')
        elif section_name == 'Inclusions':
            lines.extend(f'- {item}' for item in model['inclusions'])
        elif section_name == 'Exclusions':
            lines.extend(f'- {item}' for item in model['exclusions'])
        elif section_name == 'Test Environments':
            lines.extend(['| Name | Environment URL / Notes |', '| --- | --- |'])
            lines.extend(f'| {left} | {right} |' for left, right in model['environment_rows'])
        elif section_name == 'Defect Reporting Procedure':
            lines.extend(['| Defect Process | POC |', '| --- | --- |'])
            lines.extend(f'| {left} | {right} |' for left, right in model['defect_rows'])
        elif section_name == 'Test Strategy':
            lines.extend(f'- {item}' for item in model['strategy'])
        elif section_name == 'Test Schedule':
            lines.extend(['| Task | Dates |', '| --- | --- |'])
            lines.extend(f'| {left} | {right} |' for left, right in model['schedule_rows'])
        elif section_name == 'Test Deliverables':
            lines.extend(f'- {item}' for item in model['deliverables'])
        elif section_name == 'Entry and Exit Criteria':
            for key, values in model['entry_exit'].items():
                lines.append(f'### {key}:')
                lines.extend(f'- {item}' for item in values)
        elif section_name == 'Test Execution':
            for key, values in model['execution_entry_exit'].items():
                lines.append(f'### {key}:')
                lines.extend(f'- {item}' for item in values)
        elif section_name == 'Test Closure':
            for key, values in model['closure_entry_exit'].items():
                lines.append(f'### {key}:')
                lines.extend(f'- {item}' for item in values)
        elif section_name == 'Tools':
            lines.extend(f'- {item}' for item in model['tools'])
        elif section_name == 'Risks and Mitigations':
            lines.extend(f'- {item}' for item in model['risks'])
        elif section_name == 'Approvals':
            lines.extend(['| Approver | Signature / Date |', '| --- | --- |'])
            lines.extend(f'| {left} | {right} |' for left, right in model['approvals'])
    return '\n'.join(lines)


def generate_test_plan(prd_path: Path, output_path: Path) -> None:
    lines = extract_lines_from_docx(prd_path)
    sections = build_section_map(lines)
    write_test_plan_docx(sections, output_path)
    print(f'Generated test plan: {output_path}')


def main() -> None:
    parser = argparse.ArgumentParser(description='Generate a test plan DOCX from a PRD DOCX file.')
    parser.add_argument('--input', '-i', required=True, help='Path to the PRD .docx file')
    parser.add_argument('--output', '-o', required=False, help='Output test plan .docx file path')
    args = parser.parse_args()

    prd_path = Path(args.input)
    if not prd_path.exists() or prd_path.suffix.lower() != '.docx':
        raise ValueError('Input file must be an existing .docx file')

    output_path = Path(args.output) if args.output else prd_path.parent / f'Generated Test Plan - {prd_path.stem}.docx'
    generate_test_plan(prd_path, output_path)


if __name__ == '__main__':
    main()
